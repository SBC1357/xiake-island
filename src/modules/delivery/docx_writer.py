"""
DOCX Writer Helper

SP-6 Batch 6D: 生成 Python 产生的 docx 文件作为正式输出。

字数计算规则（仅计算正文）：
- 包含：intro/body 段落、正文中的标题、列表文本
- 排除：title/元数据/页眉页脚/表格/图表标题/脚注/附录/参考文献/系统元数据
"""
import re
from pathlib import Path
from typing import List, Dict, Any, Optional
from datetime import datetime
from dataclasses import dataclass

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


@dataclass
class DocxGenerationResult:
    """DOCX 生成结果"""
    docx_path: Path
    body_word_count: int


class DocxWriter:
    """
    DOCX 文件写入器
    
    SP-6 Batch 6D: 生成正式的 docx 文件，并计算正文字数。
    """
    
    def __init__(self, output_dir: Path = None):
        self.output_dir = output_dir or Path("output")
    
    def write(
        self,
        thesis: str,
        outline: List[Dict[str, Any]],
        key_evidence: Optional[List[str]] = None,
        content: Optional[str] = None,
        target_audience: Optional[str] = None,
        play_id: Optional[str] = None,
        arc_id: Optional[str] = None,
        filename: Optional[str] = None
    ) -> DocxGenerationResult:
        """
        写入 DOCX 文件
        
        Args:
            thesis: 核心论点
            outline: 文章大纲
            key_evidence: 核心证据
            content: 正文内容
            target_audience: 目标受众
            play_id: 写作策略ID
            arc_id: 叙事弧线ID
            filename: 自定义文件名
            
        Returns:
            DocxGenerationResult: 生成结果，包含路径和正文字数
        """
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        # 生成文件名
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            safe_title = self._safe_filename(thesis[:30])
            filename = f"{safe_title}_{timestamp}.docx"
        
        # 确保 .docx 扩展名
        if not filename.endswith(".docx"):
            filename = filename.rsplit(".", 1)[0] + ".docx"
        
        docx_path = self.output_dir / filename
        
        # 创建文档
        doc = Document()
        
        # 添加标题（不计入正文字数）
        title = doc.add_heading(thesis, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加元数据段落（不计入正文字数）
        meta_para = doc.add_paragraph()
        meta_para.add_run(f"生成时间: {datetime.now().isoformat()}")
        if target_audience:
            meta_para.add_run(f"\n目标受众: {target_audience}")
        if play_id:
            meta_para.add_run(f"\n策略: {play_id}")
        if arc_id:
            meta_para.add_run(f"\n弧线: {arc_id}")
        
        # 添加分隔
        doc.add_paragraph()
        
        # === 正文部分开始 ===
        # 正文内容用于字数计算
        
        body_text_parts = []
        
        # 生成正文内容
        if content:
            # 使用提供的内容
            body_paragraphs = self._parse_content_to_paragraphs(doc, content)
            body_text_parts.extend(body_paragraphs)
        else:
            # 从 thesis/outline/key_evidence 生成正文
            generated_paragraphs = self._generate_body_from_inputs(
                doc, thesis, outline, key_evidence
            )
            body_text_parts.extend(generated_paragraphs)
        
        # 计算正文字数
        body_word_count = self._count_body_words(body_text_parts)
        
        # 保存文档
        doc.save(str(docx_path))
        
        return DocxGenerationResult(
            docx_path=docx_path,
            body_word_count=body_word_count
        )
    
    def _safe_filename(self, text: str) -> str:
        """生成安全的文件名"""
        # 替换不安全字符
        safe = re.sub(r'[<>:"/\\|?*]', '_', text)
        return safe.strip()
    
    def _parse_content_to_paragraphs(
        self,
        doc: Document,
        content: str
    ) -> List[str]:
        """
        解析内容为段落并添加到文档
        
        Returns:
            正文文本列表
        """
        body_parts = []
        
        # 按段落分割
        paragraphs = content.split("\n\n")
        
        for para_text in paragraphs:
            para_text = para_text.strip()
            if not para_text:
                continue
            
            # 检测标题
            if para_text.startswith("# "):
                # 一级标题（不计入正文字数）
                doc.add_heading(para_text[2:], level=1)
            elif para_text.startswith("## "):
                # 二级标题
                heading_text = para_text[3:]
                doc.add_heading(heading_text, level=2)
                body_parts.append(heading_text)
            elif para_text.startswith("### "):
                # 三级标题
                heading_text = para_text[4:]
                doc.add_heading(heading_text, level=3)
                body_parts.append(heading_text)
            else:
                # 普通段落
                # 检测列表
                if para_text.startswith("- ") or para_text.startswith("* "):
                    # 无序列表
                    lines = para_text.split("\n")
                    for line in lines:
                        line = line.strip()
                        if line.startswith("- ") or line.startswith("* "):
                            list_text = line[2:]
                            doc.add_paragraph(list_text, style='List Bullet')
                            body_parts.append(list_text)
                else:
                    # 普通段落
                    doc.add_paragraph(para_text)
                    body_parts.append(para_text)
        
        return body_parts
    
    def _generate_body_from_inputs(
        self,
        doc: Document,
        thesis: str,
        outline: List[Dict[str, Any]],
        key_evidence: Optional[List[str]] = None
    ) -> List[str]:
        """
        从 thesis/outline/key_evidence 生成正文段落
        
        当 content 缺失时，生成确定性的正文内容。
        
        Returns:
            正文文本列表
        """
        body_parts = []
        
        # 添加引言
        doc.add_heading("引言", level=2)
        body_parts.append("引言")
        
        intro_text = f"本文旨在探讨{thesis}。"
        doc.add_paragraph(intro_text)
        body_parts.append(intro_text)
        
        # 添加大纲章节
        if outline:
            doc.add_heading("主要内容", level=2)
            body_parts.append("主要内容")
            
            for i, section in enumerate(outline, 1):
                title = section.get("title", f"第{i}节")
                
                # 添加章节标题
                doc.add_heading(title, level=3)
                body_parts.append(title)
                
                # 生成章节内容
                section_type = section.get("type", "section")
                
                if section_type == "domain_section":
                    fact_count = section.get("fact_count", 0)
                    domain = section.get("domain", "相关")
                    para_text = f"在{domain}方面，本节基于{fact_count}条证据进行分析。"
                elif section_type == "evidence":
                    fact_id = section.get("fact_id", "")
                    para_text = f"本节聚焦于证据{fact_id}的深入分析。"
                elif section_type == "intro":
                    para_text = f"{title}部分将为全文奠定基础。"
                elif section_type == "conclusion":
                    para_text = f"{title}部分将总结全文核心观点。"
                else:
                    para_text = f"本节讨论{title}相关内容。"
                
                doc.add_paragraph(para_text)
                body_parts.append(para_text)
        
        # 添加核心证据
        if key_evidence:
            doc.add_heading("核心证据", level=2)
            body_parts.append("核心证据")
            
            for evidence in key_evidence:
                doc.add_paragraph(evidence, style='List Bullet')
                body_parts.append(evidence)
        
        # 添加结论
        doc.add_heading("结论", level=2)
        body_parts.append("结论")
        
        conclusion_text = f"综上所述，{thesis}。本文通过系统性分析，为相关决策提供了参考依据。"
        doc.add_paragraph(conclusion_text)
        body_parts.append(conclusion_text)
        
        return body_parts
    
    def _count_body_words(self, text_parts: List[str]) -> int:
        """
        计算正文字数
        
        中文字数按字符计，英文按单词计。
        """
        total_chars = 0
        
        for text in text_parts:
            # 移除多余空白
            text = text.strip()
            if not text:
                continue
            
            # 统计中文字符
            chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', text))
            
            # 统计英文单词（简化处理）
            english_words = len(re.findall(r'[a-zA-Z]+', text))
            
            # 统计数字
            numbers = len(re.findall(r'\d+', text))
            
            # 总字数 = 中文 + 英文单词 + 数字
            total_chars += chinese_chars + english_words + numbers
        
        return total_chars